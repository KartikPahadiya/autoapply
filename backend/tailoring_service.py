"""
Resume tailoring + cover letter generation using CV Forge MCP server.

Flow:
1. Parse resume text into structured userProfile (using NVIDIA LLM)
2. Call cv-forge's draft_complete_application — generates CV PDF + cover letter + email
3. Read the generated files from disk
4. Return PDF bytes + cover letter text

Requires Node.js 18+ installed. The MCP server auto-downloads via npx.
"""
import base64
import io
import json
import os
import re
import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_nvidia_ai_endpoints import ChatNVIDIA
from langgraph.prebuilt import create_react_agent
from pypdf import PdfReader

from mcp_tools import get_mcp_tools

NVIDIA_API_KEY = os.getenv("NVIDIA_API_KEY")

_nvidia_llm: ChatNVIDIA | None = None
_tailoring_agent = None


def nvidia_configured() -> bool:
    return bool(NVIDIA_API_KEY)


def _get_nvidia_llm() -> ChatNVIDIA:
    global _nvidia_llm
    if _nvidia_llm is None:
        if not nvidia_configured():
            raise RuntimeError("NVIDIA_API_KEY is not set in .env — resume tailoring is unavailable.")
        _nvidia_llm = ChatNVIDIA(
            model="nvidia/nemotron-3-ultra-550b-a55b",
            api_key=NVIDIA_API_KEY,
            temperature=0.7,
            top_p=0.95,
            max_tokens=4096,
        )
    return _nvidia_llm


async def _get_tailoring_agent():
    global _tailoring_agent
    if _tailoring_agent is None:
        tools = await get_mcp_tools()
        _tailoring_agent = create_react_agent(_get_nvidia_llm(), tools)
    return _tailoring_agent


RESUME_PARSE_PROMPT = """You are a resume parser. Extract structured information from the resume text below and return ONLY a JSON object with these fields:

{
  "fullName": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedIn": "",
  "summary": "",
  "experience": "",
  "education": "",
  "skills": "",
  "certifications": ""
}

Return ONLY the JSON object, no markdown, no explanation. Use empty strings for missing fields. Keep each field as a single string (concatenate multiple items with commas or newlines as needed)."""


TAILORING_SYSTEM_PROMPT = """You are a resume tailoring assistant with access to the CV Forge MCP tools.

You have these tools available:
- draft_complete_application(userProfile, jobRequirements, outputPath, baseFileName) — generates a complete job application package: CV PDF, cover letter PDF, and email template. Saves files to the outputPath directory.
- generate_cv(userProfile, jobRequirements, outputPath, fileName, format) — generates a tailored CV. Defaults to PDF.
- generate_cover_letter(userProfile, jobRequirements, hiringManagerName) — generates a cover letter.

Your task:
1. Call draft_complete_application with the user's profile, job requirements, and a temp output directory. Use a baseFileName like "tailored".
2. The tool will save files to the outputPath. Report what files were created.

If draft_complete_application is not available, fall back to generate_cv + generate_cover_letter separately."""


async def _parse_resume_to_profile(resume_text: str) -> dict:
    """Use NVIDIA LLM to extract structured profile from raw resume text."""
    llm = _get_nvidia_llm()
    prompt = f"{RESUME_PARSE_PROMPT}\n\nResume text:\n{resume_text}\n\nJSON:"
    msg = await llm.ainvoke([HumanMessage(content=prompt)])
    raw = msg.content.strip()

    # Extract JSON block
    json_match = re.search(r"\{[\s\S]*?\}", raw)
    if json_match:
        try:
            return json.loads(json_match.group(0))
        except json.JSONDecodeError:
            pass

    # Fallback: return minimal profile with raw text
    return {
        "fullName": "Applicant",
        "email": "",
        "phone": "",
        "location": "",
        "linkedIn": "",
        "summary": "",
        "experience": resume_text[:2000],
        "education": "",
        "skills": "",
        "certifications": "",
    }


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def text_to_docx_bytes(title: str, text: str) -> bytes:
    document = DocxDocument()
    if title:
        document.add_heading(title, level=1)
    for block in re.split(r"\n\s*\n", text.strip()):
        cleaned = block.strip()
        if cleaned:
            document.add_paragraph(cleaned)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


async def tailor_resume_and_cover_letter(
    resume_text: str, job_description: str, company: str = "", title: str = ""
) -> dict:
    """Runs the CV Forge MCP pipeline.

    Returns {"cover_letter": str, "tailored_resume_url": str | None,
    "tailored_resume_text": str, "raw_reply": str}.
    """
    # 1. Parse resume into structured profile
    user_profile = await _parse_resume_to_profile(resume_text)

    # 2. Create temp output directory
    output_dir = tempfile.mkdtemp(prefix="cv_forge_")
    base_name = "tailored"

    # 3. Build job requirements object
    job_requirements = {
        "jobTitle": title or "unspecified",
        "company": company or "unspecified",
        "jobDescription": job_description,
    }

    # 4. Run the agent with CV Forge tools
    agent = await _get_tailoring_agent()

    user_prompt = (
        f"User Profile:\n{json.dumps(user_profile, indent=2)}\n\n"
        f"Job Requirements:\n{json.dumps(job_requirements, indent=2)}\n\n"
        f"Call draft_complete_application with:\n"
        f"- userProfile: the user profile above\n"
        f"- jobRequirements: the job requirements above\n"
        f"- outputPath: {output_dir}\n"
        f"- baseFileName: {base_name}\n\n"
        f"If draft_complete_application is not available, use generate_cv + generate_cover_letter separately."
    )

    result = await agent.ainvoke(
        {"messages": [SystemMessage(content=TAILORING_SYSTEM_PROMPT), HumanMessage(content=user_prompt)]}
    )
    messages = result["messages"]
    final_reply = messages[-1].content if messages else ""

    # 5. Try to read generated files from the output directory
    pdf_bytes = None
    docx_bytes = None
    cover_letter_text = None
    tailored_resume_text = None

    expected_pdf = Path(output_dir) / f"{base_name}_CV.pdf"
    expected_cover = Path(output_dir) / f"{base_name}_Cover_Letter.pdf"
    expected_email = Path(output_dir) / f"{base_name}_Email_Template.txt"

    if expected_pdf.exists():
        pdf_bytes = expected_pdf.read_bytes()
        tailored_resume_text = _extract_pdf_text(pdf_bytes)

    if expected_cover.exists():
        cover_letter_text = _extract_pdf_text(expected_cover.read_bytes())

    if expected_email.exists():
        email_text = expected_email.read_text(encoding="utf-8", errors="replace")
        if not cover_letter_text:
            cover_letter_text = email_text

    # 6. Fallback: extract from tool messages if files weren't found
    if not cover_letter_text:
        for msg in messages:
            content = str(getattr(msg, "content", ""))
            if len(content) > 200 and not content.startswith("{"):
                cover_letter_text = content
                break

    docx_source = tailored_resume_text or resume_text
    if docx_source:
        docx_bytes = text_to_docx_bytes("Tailored Resume", docx_source)

    # 7. Build return value
    return {
        "cover_letter": cover_letter_text or "Cover letter generation failed.",
        "tailored_resume_url": None,  # We have bytes, not a URL
        "tailored_resume_text": tailored_resume_text or "",
        "pdf_bytes": pdf_bytes,  # New field: raw PDF bytes
        "docx_bytes": docx_bytes,
        "docx_filename": "tailored_resume.docx",
        "raw_reply": final_reply,
    }


def download_pdf(url: str) -> bytes | None:
    """No-op — CV Forge saves files locally. Kept for API compatibility."""
    return None
